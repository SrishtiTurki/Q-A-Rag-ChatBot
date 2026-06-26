import fitz  # PyMuPDF
import docx
import csv
import io
import os
from PIL import Image
from pathlib import Path
from ingestion.ocr import extract_text_from_image


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".png", ".jpg", ".jpeg", ".webp", ".tiff"}


def _is_clean_ocr(text: str) -> bool:
    if not text or len(text.strip()) < 30:
        return False
    words = text.split()
    if len(words) < 8:
        return False
    single_chars = sum(1 for w in words if len(w) == 1)
    if single_chars / len(words) > 0.4:
        return False
    noise_patterns = ["INFO:", "HTTP/1.1", "127.0.0.1", "Traceback", "WARNING:", "ERROR:"]
    if any(p in text for p in noise_patterns):
        return False
    return True


def parse_file(file_path: str) -> list[dict]:
    """Full parse — text only, no OCR. Returns page dicts."""
    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")
    if ext == ".pdf":
        return parse_pdf_text(file_path)
    elif ext == ".docx":
        return parse_docx_text(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    elif ext == ".csv":
        return parse_csv(file_path)
    elif ext in {".png", ".jpg", ".jpeg", ".webp", ".tiff"}:
        # standalone image — goes straight to OCR
        text = extract_text_from_image(file_path)
        return [{"text": text, "page_number": 1, "line_count": text.count('\n') + 1 if text else 0}]
    return []


def parse_images_only(file_path: str) -> list[dict]:
    """
    Phase 2 — extract and OCR embedded images only.
    Called in background after text is already indexed.
    Returns page dicts same format as text parser.
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        pages = _extract_pdf_images(file_path)
        print(f"[Image Parser] Extracted {len(pages)} image pages from PDF")
        return pages
    elif ext == ".docx":
        pages = _extract_docx_images(file_path)
        print(f"[Image Parser] Extracted {len(pages)} image pages from DOCX")
        return pages
    return []  # TXT/CSV have no images


def _extract_pdf_images(file_path: str) -> list[dict]:
    """Extract and OCR embedded images from PDF. Skips small/decorative images."""
    doc = fitz.open(file_path)
    pages = []
    image_count = 0
    
    for page_num, page in enumerate(doc):
        image_list = page.get_images(full=True)
        if image_list:
            print(f"[Image Parser] Page {page_num + 1} has {len(image_list)} images")
        
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            base_image = doc.extract_image(xref)

            # Skip small images — likely logos or decorations
            width = base_image.get("width", 0)
            height = base_image.get("height", 0)
            if width < 150 or height < 150:
                continue

            try:
                img = Image.open(io.BytesIO(base_image["image"])).convert("RGB")
                img_text = extract_text_from_image(img)
                
                if img_text.strip() and _is_clean_ocr(img_text):
                    image_count += 1
                    pages.append({
                        "text": f"[Image {image_count} on Page {page_num + 1}]\n{img_text}",
                        "page_number": page_num + 1,
                        "line_count": img_text.count('\n') + 1 if img_text else 0
                    })
                    print(f"[Image Parser] Extracted text from image {image_count} on page {page_num + 1}")
            except Exception as e:
                print(f"[Image Parser] Error extracting image on page {page_num + 1}: {e}")
                continue
                
    doc.close()
    print(f"[Image Parser] Total: {len(pages)} images extracted with text")
    return pages


# ─── Text parsers ─────────────────────────────────────────────────────────────

def parse_pdf_text(file_path: str) -> list[dict]:
    """Extract selectable text from PDF, page by page. No OCR."""
    doc = fitz.open(file_path)
    pages = []
    for page_num, page in enumerate(doc):
        page_text = page.get_text("text").strip()

        # Only OCR if page has zero selectable text (fully scanned page)
        if not page_text:
            pix = page.get_pixmap(dpi=150)  # lower DPI = faster
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_text = extract_text_from_image(img)

        if page_text.strip():
            pages.append({
                "text": page_text,
                "page_number": page_num + 1,
                "line_count": page_text.count('\n') + 1 if page_text else 0
            })
    doc.close()
    return pages


def parse_docx_text(file_path: str) -> list[dict]:
    """Extract text paragraphs and tables from DOCX. No image OCR."""
    document = docx.Document(file_path)
    all_paragraphs = []

    for para in document.paragraphs:
        if para.text.strip():
            all_paragraphs.append(para.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                all_paragraphs.append(row_text)

    pages = []
    page_size = 50
    for i in range(0, len(all_paragraphs), page_size):
        group = all_paragraphs[i:i + page_size]
        full_text = "\n\n".join(group)
        pages.append({
            "text": full_text,
            "page_number": (i // page_size) + 1,
            "line_count": full_text.count('\n') + 1 if full_text else 0
        })
    return pages


def parse_txt(file_path: str) -> list[dict]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        lines = f.readlines()
    pages = []
    page_size = 100
    for i in range(0, len(lines), page_size):
        full_text = "".join(lines[i:i + page_size]).strip()
        pages.append({
            "text": full_text,
            "page_number": (i // page_size) + 1,
            "line_count": len(lines[i:i + page_size])
        })
    return pages


def parse_csv(file_path: str) -> list[dict]:
    rows_text = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row_str = ", ".join(f"{k}: {v}" for k, v in row.items() if v.strip())
            if row_str:
                rows_text.append(row_str)
    pages = []
    page_size = 50
    for i in range(0, len(rows_text), page_size):
        full_text = "\n".join(rows_text[i:i + page_size])
        pages.append({
            "text": full_text,
            "page_number": (i // page_size) + 1,
            "line_count": len(rows_text[i:i + page_size])
        })
    return pages


# ─── Image extractors (Phase 2) ───────────────────────────────────────────────

'''def _extract_pdf_images(file_path: str) -> list[dict]:
    """Extract and OCR embedded images from PDF. Skips small/decorative images."""
    doc = fitz.open(file_path)
    pages = []
    for page_num, page in enumerate(doc):
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            base_image = doc.extract_image(xref)

            # Skip small images — likely logos or decorations
            if base_image.get("width", 0) < 150 or base_image.get("height", 0) < 150:
                continue

            img = Image.open(io.BytesIO(base_image["image"])).convert("RGB")
            img_text = extract_text_from_image(img)
            if img_text.strip() and _is_clean_ocr(img_text):
                pages.append({
                    "text": f"[Image on Page {page_num + 1}]\n{img_text}",
                    "page_number": page_num + 1,
                    "line_count": img_text.count('\n') + 1 if img_text else 0
                })
    doc.close()
    return pages

def _extract_docx_images(file_path: str) -> list[dict]:
    """Extract and OCR embedded images from DOCX."""
    document = docx.Document(file_path)
    pages = []
    image_count = 0
    
    for idx, rel in enumerate(document.part.rels.values()):
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                img = Image.open(io.BytesIO(img_data)).convert("RGB")

                # Skip small images
                if img.width < 150 or img.height < 150:
                    continue

                img_text = extract_text_from_image(img)
                if img_text.strip() and _is_clean_ocr(img_text):
                    image_count += 1
                    pages.append({
                        "text": f"[Image {image_count}]\n{img_text}",
                        "page_number": len(pages) + 1,
                        "line_count": img_text.count('\n') + 1 if img_text else 0
                    })
                    print(f"[Image Parser] Extracted text from image {image_count}")
            except Exception as e:
                print(f"[Image Parser] Error extracting image: {e}")
                continue
                
    print(f"[Image Parser] Total: {len(pages)} images extracted with text")
    return pages
'''